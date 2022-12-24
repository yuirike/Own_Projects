#Main file, the "code" is here.

import docx
class VN_Summary_Extractor:
    def __init__(self, file):
        self.__file = file

    @staticmethod
    def __filter(data):
        for x in data:
            if ":" in x:
                position = data.index(x)
                new_x = x[x.find(":")+1:]
                data[position] = new_x
        return data

    @staticmethod
    def data_processor(self, data):
        #Extracting numbers based on predetermined format.
        data = [x for x in data if ";" or "h" in x]
        minutes = [x for x in data if ";" in x]
        hours = [x for x in data if "h" in x]

        #Filtering
        hours = [x[:x.find("h")] for x in hours]
        minutes = [x[:x.find(";")] for x in minutes]
        hours = self.__filter(hours)
        minutes = self.__filter(minutes)

        #Math
        hours = [int(x) for x in hours]
        minutes = [int(x) for x in minutes]
        min_hour = (sum(minutes)- (sum(minutes)%60))//60
        min = sum(minutes)%60
        h = sum(hours) + min_hour
        return "Total Time was {}h {}min".format(h, min)

    def time(self):
        data = []
        doc = docx.Document(self.__file)
        for line in doc.paragraphs[0:5]:
            data.append(line.text)

        #Setting up data
        data = [x for x in data if "Time" in x]
        data = data[0].split()
        return self.data_processor(self, data),  data

    def old_time(self):
        from datetime import datetime
        data = []
        doc = docx.Document(self.__file)
        for line in doc.paragraphs[0:5]:
            data.append(line.text)
        
        return data

        


